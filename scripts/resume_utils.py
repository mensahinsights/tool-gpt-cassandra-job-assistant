#!/usr/bin/env python3
import json
from docx import Document

def load_baseline(path="data/baseline_resume.json"):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def render_resume(output_path: str, baseline: dict, tailoring: dict):
    doc = Document()

    # Header
    doc.add_heading(baseline["header"]["name"], 0)
    doc.add_paragraph(f"{baseline['header']['city']} | {baseline['header']['email']}")
    doc.add_paragraph(" | ".join(baseline["header"]["links"]))

    # Headline
    doc.add_heading("Headline", level=1)
    doc.add_paragraph(tailoring.get("headline", baseline["headline"]))

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(tailoring.get("summary", baseline["summary"]))

    # Core Skills
    doc.add_heading("Core Skills", level=1)
    skills = tailoring.get("core_skills", baseline["core_skills"])
    doc.add_paragraph(", ".join(skills))

    # Experience
    doc.add_heading("Experience", level=1)
    for role in baseline["experience"]:
        title_line = f"{role['title']} | {role['company']} | {role['location']} | {role['start_date']} – {role['end_date']}"
        doc.add_paragraph(title_line, style="Intense Quote")
        doc.add_paragraph(role["impact"])
        for b in role["bullets"]:
            doc.add_paragraph(f"• {b}", style="List Bullet")

    # Education
    doc.add_heading("Education & Certifications", level=1)
    for ed in baseline["education"]:
        line = f"{ed.get('degree','')} – {ed.get('field','')} ({ed.get('institution','')})"
        doc.add_paragraph(line)

    # ATS Score
    doc.add_heading("ATS Score", level=1)
    doc.add_paragraph(str(tailoring.get("ats_score", "")))

    doc.save(output_path)
