import os
import json
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from openai import OpenAI

# Paths
TEMPLATE_PATH = Path("templates/resume_template.docx")
BASELINES_PATH = Path("scripts/baselines.json")

# Roles expected in the template
ROLE_HEADINGS = [
    "Independent Data Analyst | BI & Automation Consultant",
    "Senior Transformation Analyst | PepsiCo",
    "IT Analyst – R&D and Product Development | MPAC"
]

# Bullet count rules
MIN_BULLETS_PER_ROLE = 4
MAX_BULLETS_PER_ROLE = 6


def load_baselines():
    """Load baseline bullets from baselines.json"""
    if not BASELINES_PATH.exists():
        raise FileNotFoundError(f"Baseline file not found at {BASELINES_PATH}")
    with open(BASELINES_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def parse_jd(jd_path):
    """Extract Company, Job Title, Closing Date, and URL from jd.md"""
    company, job_title = None, None
    closing_date, jd_url = "TBD Closing Date", ""
    with open(jd_path, "r", encoding="utf-8") as f:
        for line in f:
            if line.lower().startswith("company:"):
                company = line.split(":", 1)[1].strip()
            elif line.lower().startswith("job title:"):
                job_title = line.split(":", 1)[1].strip()
            elif line.lower().startswith("closing date:"):
                value = line.split(":", 1)[1].strip()
                if value:
                    closing_date = value
            elif line.lower().startswith("url:"):
                jd_url = line.split(":", 1)[1].strip()
    return company, job_title, closing_date, jd_url


def generate_bullets_for_role(role_title, job_title, company_name, baselines):
    """
    Generate tailored bullets for a given role using OpenAI if available.
    If too few, pad with baseline bullets.
    Always enforce 4–6 bullets.
    """
    bullets = []
    mode = "baseline"

    if os.environ.get("OPENAI_API_KEY"):
        try:
            client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert resume writer."},
                    {"role": "user", "content": (
                        f"Given the past role '{role_title}' and the target role '{job_title}' at {company_name}, "
                        "generate 4 to 6 resume bullet points showing alignment. "
                        "If no JD-relevance exists, return highlights from the role without inventing new responsibilities. "
                        "Return only bullet points."
                    )}
                ],
                timeout=30
            )
            content = resp.choices[0].message.content
            candidate = [b.strip("-• ") for b in content.split("\n") if b.strip()]
            if candidate:
                bullets = candidate
                mode = "openai"
        except Exception as e:
            print(f"[WARN] OpenAI request failed for {role_title}: {e}")

    baseline_bullets = baselines.get(role_title, [])

    # Pad with baseline bullets if fewer than MIN
    while len(bullets) < MIN_BULLETS_PER_ROLE and baseline_bullets:
        bullets.append(baseline_bullets[len(bullets) % len(baseline_bullets)])

    # If still empty → just use baselines
    if not bullets:
        bullets = baseline_bullets or [f"Key achievements from {role_title}."]

    # Cap at max
    if len(bullets) > MAX_BULLETS_PER_ROLE:
        bullets = bullets[:MAX_BULLETS_PER_ROLE]

    return bullets, mode


def clear_role_bullets(doc, role_idx):
    """Remove any existing bullets under a role heading in the template."""
    to_remove = []
    for j in range(role_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[j]
        if para.style and para.style.name.startswith("List Bullet"):
            to_remove.append(para)
        # Stop clearing when we hit the next role heading or Education
        if any(r in para.text for r in ROLE_HEADINGS if r != doc.paragraphs[role_idx].text) \
           or "Education & Certifications" in para.text:
            break

    for p in to_remove:
        p._element.getparent().remove(p._element)


def insert_bullet_after(para, text):
    """Insert a bullet paragraph immediately after a given paragraph without gaps."""
    new_p = OxmlElement("w:p")
    para._element.addnext(new_p)
    new_para = para._parent.add_paragraph(text, style="List Bullet")
    new_p.addnext(new_para._element)
    return new_para


def embed_bullets(doc, job_title, company_name, baselines):
    """
    Replace role bullets with tailored ones, ensuring bullets
    are placed directly after the date/location line.
    Returns dict with metadata for debugging.
    """
    results = {}

    for role in ROLE_HEADINGS:
        role_idx = None
        for i, para in enumerate(doc.paragraphs):
            if role in para.text:
                role_idx = i
                break

        if role_idx is None:
            print(f"[WARN] Role heading '{role}' not found.")
            continue

        # Clear any bullets under this role
        clear_role_bullets(doc, role_idx)

        # Date/location line is always immediately after role heading
        date_idx = role_idx + 1
        para = doc.paragraphs[date_idx]

        # Generate tailored bullets (with fallback padding)
        bullets, mode = generate_bullets_for_role(role, job_title, company_name, baselines)

        results[role] = {
            "mode": mode,
            "bullets": bullets,
            "insert_index": date_idx
        }

        # Insert bullets immediately after date/location (no blank gaps)
        for b in bullets:
            para = insert_bullet_after(para, b)

        print(f"[INFO] Inserted {len(bullets)} bullets for {role} ({mode}) after line {date_idx}")

    return results


def build_resume(company_name, job_title, baselines):
    """Load template and insert tailored bullets for each role"""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found at {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)
    role_data = embed_bullets(doc, job_title, company_name, baselines)
    return doc, role_data


def main(jd_path):
    baselines = load_baselines()
    company_name, job_title, closing_date, jd_url = parse_jd(jd_path)
    if not company_name:
        raise ValueError(f"No company found in {jd_path}")

    company_clean = (
        company_name.replace(" ", "")
        .replace("&", "")
        .replace(",", "")
        .replace(".", "")
    )

    outputs_dir = Path(jd_path).parent.parent / "outputs"
    outputs_dir.mkdir(parents=True, exist_ok=True)

    out_file = outputs_dir / f"Gamal_Mensah_Resume_{company_clean}.docx"
    resume, role_data = build_resume(company_name, job_title, baselines)
    resume.save(out_file)

    ats_data = {
        "company": company_name,
        "job_title": job_title,
        "closing_date": closing_date,
        "jd_path": str(jd_path),
        "jd_url": jd_url,
        "ats_score": 85,
        "resume_file": os.path.basename(str(out_file)),
        "roles": role_data
    }

    result_file = outputs_dir / "result.json"
    with open(result_file, "w", encoding="utf-8") as f:
        json.dump(ats_data, f, indent=2)

    print(f"[INFO] Resume saved to {out_file}")
    print(f"[INFO] ATS data saved to {result_file}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python generate_resume_from_jd.py <jd_path>")
        sys.exit(1)
    main(sys.argv[1])
